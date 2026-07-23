"""
docgen.py — Warehouse pick list & delivery note generator.

Reads the daily produce export (same layout as the example file) and produces:
  • Pick list  : one Word file per VEHICLE  (shops broken down within each file)
  • Delivery   : one Word file per STORE

Only 5 source columns are used (by position, robust to header text changes):
  A (0) Vehicle | C (2) Store | E (4) Barcode | F (5) Item | I (8) Boxes
"""

import io
import re
import zipfile

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ----------------------------------------------------------------------------
# Column positions in the source file (0-based). "Exactly like the example."
# ----------------------------------------------------------------------------
COL_VEHICLE  = 0   # A
COL_STORE    = 2   # C
COL_PRIORITY = 3   # D  (item/species priority — orders SKUs on delivery notes)
COL_BARCODE  = 4   # E
COL_ITEM     = 5   # F
COL_BOXES    = 8   # I
MIN_COLUMNS  = 9

FONT = "Arial"
C_HEADER_FILL = "1F3864"
C_ALT_FILL    = "F2F2F2"
C_TOTAL_FILL  = "DDEBF7"
C_SHOP_FILL   = "E8EEF7"
C_GREY        = "595959"
C_ACCENT      = "1F3864"

ALIGN = {
    "left":   WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right":  WD_ALIGN_PARAGRAPH.RIGHT,
}

# column widths (inches) — sum ≈ 7.5" content width (Letter, 0.5" margins)
PICK_WIDTHS = [0.36, 3.81, 1.67, 0.83, 0.83]          # №, item, barcode, boxes, ✓
DEL_WIDTHS  = [0.39, 4.12, 1.49, 0.75, 0.75]          # №, item, barcode, boxes, ✓


# ============================================================================
# 1. DATA LOADING & AGGREGATION
# ============================================================================
def load_and_prepare(file_like):
    """Read the export, keep the 5 used columns, clean, and return a tidy frame.

    Returns (df, stats) where df has columns:
        Vehicle, Store, Barcode, Item, Boxes
    and stats is a dict with row counts for a data-quality summary.
    """
    raw = pd.read_excel(file_like, header=None, skiprows=1, dtype=str)
    if raw.shape[1] < MIN_COLUMNS:
        raise ValueError(
            f"ფაილში მოსალოდნელია მინიმუმ {MIN_COLUMNS} სვეტი (A–I), "
            f"ნაპოვნია {raw.shape[1]}. შეამოწმეთ, რომ ფაილი ნიმუშის იდენტურია."
        )

    df = pd.DataFrame({
        "Vehicle":  raw.iloc[:, COL_VEHICLE],
        "Store":    raw.iloc[:, COL_STORE],
        "Priority": raw.iloc[:, COL_PRIORITY],
        "Barcode":  raw.iloc[:, COL_BARCODE],
        "Item":     raw.iloc[:, COL_ITEM],
        "Boxes":    raw.iloc[:, COL_BOXES],
    })

    total_read = len(df)
    for c in ["Vehicle", "Store", "Barcode", "Item"]:
        df[c] = df[c].fillna("").astype(str).str.strip()
    df["Boxes"] = pd.to_numeric(df["Boxes"], errors="coerce")
    df["Priority"] = pd.to_numeric(df["Priority"], errors="coerce")

    # drop unusable rows (a missing Priority is tolerated — see below)
    bad_mask = (
        (df["Vehicle"] == "") | (df["Store"] == "") |
        (df["Barcode"] == "") | (df["Item"] == "") |
        df["Boxes"].isna() | (df["Boxes"] <= 0)
    )
    dropped = int(bad_mask.sum())
    df = df[~bad_mask].copy()
    df["Boxes"] = df["Boxes"].round().astype(int)
    # items without a priority value sort to the end, not the top
    df["Priority"] = df["Priority"].fillna(10**9)

    if df.empty:
        raise ValueError("გამოსაყენებელი მონაცემები ვერ მოიძებნა (ცარიელი ან არასწორი ფაილი).")

    # defensive: collapse duplicate Vehicle×Store×Barcode; keep summed boxes
    # and the item's priority (min, since it is constant per SKU)
    df = (df.groupby(["Vehicle", "Store", "Barcode", "Item"], as_index=False, sort=False)
            .agg(Boxes=("Boxes", "sum"), Priority=("Priority", "min")))

    stats = {
        "rows_read": total_read,
        "rows_used": len(df),
        "rows_dropped": dropped,
        "vehicles": df["Vehicle"].nunique(),
        "stores": df["Store"].nunique(),
        "skus": df["Barcode"].nunique(),
        "total_boxes": int(df["Boxes"].sum()),
    }
    return df, stats


def _store_source_order(df):
    """Return list of (vehicle, store) in the order stores first appear."""
    order = (df.assign(_i=range(len(df)))
               .groupby(["Vehicle", "Store"], sort=False)["_i"].min()
               .reset_index().sort_values("_i"))
    return list(zip(order["Vehicle"], order["Store"]))


def build_pick_nested(df):
    """vehicle -> {shops -> {rows}}  (shops in source order, items alphabetical)."""
    pairs = _store_source_order(df)
    out = []
    for veh in sorted(df["Vehicle"].unique()):
        dv = df[df["Vehicle"] == veh]
        shops = []
        for (v, st) in pairs:
            if v != veh:
                continue
            ds = dv[dv["Store"] == st].sort_values("Item", kind="stable")
            rows = [{"item": r.Item, "barcode": r.Barcode, "boxes": int(r.Boxes)}
                    for r in ds.itertuples()]
            shops.append({"store": st, "boxes": int(ds["Boxes"].sum()), "rows": rows})
        out.append({
            "vehicle": veh,
            "shop_count": len(shops),
            "total_boxes": int(dv["Boxes"].sum()),
            "shops": shops,
        })
    return out


def build_vehicle_delivery(df):
    """One consolidated entry per vehicle: SKU totals across all its stores.

    This is the warehouse -> driver handover note (what the truck carries in
    total). Sorted by item priority, same convention as the store notes.
    """
    out = []
    for veh in sorted(df["Vehicle"].unique()):
        dv = df[df["Vehicle"] == veh]
        agg = (dv.groupby(["Barcode", "Item"], as_index=False)
                 .agg(Boxes=("Boxes", "sum"), Priority=("Priority", "min"))
                 .sort_values(["Priority", "Item"], kind="stable"))
        rows = [{"item": r.Item, "barcode": r.Barcode, "boxes": int(r.Boxes)}
                for r in agg.itertuples()]
        out.append({
            "vehicle": veh,
            "store_count": int(dv["Store"].nunique()),
            "boxes": int(dv["Boxes"].sum()),
            "rows": rows,
        })
    return out


def build_delivery_flat(df):
    """One entry per store (vehicle -> store order), items alphabetical."""
    pairs = _store_source_order(df)
    out = []
    for (veh, st) in pairs:
        ds = (df[(df["Vehicle"] == veh) & (df["Store"] == st)]
              .sort_values(["Priority", "Item"], kind="stable"))
        rows = [{"item": r.Item, "barcode": r.Barcode, "boxes": int(r.Boxes)}
                for r in ds.itertuples()]
        out.append({"vehicle": veh, "store": st,
                    "boxes": int(ds["Boxes"].sum()), "rows": rows})
    return out


# ============================================================================
# 2. LOW-LEVEL DOCX HELPERS
# ============================================================================
def _set_run_font(run, name=FONT, size=None, bold=None, color=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def _shade(el_pr, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    el_pr.append(shd)


def _cell_bg(cell, hex_color):
    _shade(cell._tc.get_or_add_tcPr(), hex_color)


def _para_bg(paragraph, hex_color):
    _shade(paragraph._p.get_or_add_pPr(), hex_color)


def _vcenter(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    v = OxmlElement("w:vAlign")
    v.set(qn("w:val"), "center")
    tcPr.append(v)


def _para_bottom_border(paragraph, color=C_ACCENT, sz=10):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def _table_borders(table, color="BFBFBF", sz=4):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def _fixed_layout(table, widths_in):
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)

    # overall table width
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(int(round(sum(widths_in) * 1440))))

    # rewrite the grid — LibreOffice/Word use tblGrid under fixed layout
    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for w in widths_in:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(round(w * 1440))))
        grid.append(gc)
    tblPr.addnext(grid)

    # per-cell widths too
    for row in table.rows:
        for i, w in enumerate(widths_in):
            row.cells[i].width = Inches(w)


def _repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)


def _cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    cs = OxmlElement("w:cantSplit")
    cs.set(qn("w:val"), "true")
    trPr.append(cs)


def _cell_text(cell, text, align="left", bold=False, color="000000", size=10, fill=None):
    if fill:
        _cell_bg(cell, fill)
    _vcenter(cell)
    p = cell.paragraphs[0]
    p.alignment = ALIGN[align]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("" if text is None else str(text))
    _set_run_font(r, FONT, size, bold, color)


def _setup_page(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.75)
    sec.header_distance = Inches(0.3)
    sec.bottom_margin = Inches(0.5)
    sec.left_margin = Inches(0.5)
    sec.right_margin = Inches(0.5)
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), FONT)
    pf = style.paragraph_format
    pf.space_after = Pt(2)
    pf.line_spacing = 1.0


def _page_header(doc, label, value, right_text=None, content_width=7.5):
    """Add a repeating page header: "<label> <value>" left, right_text right.

    Word renders section headers on EVERY page, so the vehicle stays visible
    even if a note overflows onto a second page.
    """
    hdr_p = doc.sections[0].header.paragraphs[0]
    hdr_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hdr_p.paragraph_format.space_after = Pt(2)
    hdr_p.paragraph_format.tab_stops.add_tab_stop(
        Inches(content_width), WD_TAB_ALIGNMENT.RIGHT
    )
    _para_bottom_border(hdr_p, color="BFBFBF", sz=6)

    r = hdr_p.add_run(label)
    _set_run_font(r, FONT, 10, False, C_GREY)
    r = hdr_p.add_run(value)
    _set_run_font(r, FONT, 11, True, C_ACCENT)      # vehicle = bold, accent
    if right_text:
        r = hdr_p.add_run("\t" + right_text)
        _set_run_font(r, FONT, 10, False, C_GREY)
    return hdr_p


def _add_para(doc, runs, space_after=4, space_before=0, keep_next=True,
              bottom_border=False, fill=None):
    """runs = list of (text, dict-of-run-kwargs)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.keep_with_next = keep_next
    if fill:
        _para_bg(p, fill)
    if bottom_border:
        _para_bottom_border(p)
    for text, kw in runs:
        r = p.add_run(text)
        _set_run_font(r, FONT, kw.get("size", 10), kw.get("bold", False), kw.get("color", "000000"))
    return p


# ============================================================================
# 3. TABLE + DOCUMENT BUILDERS
# ============================================================================
def _item_table(doc, rows, widths, with_check, total_value=None, body_size=10, check_glyph=""):
    ncols = len(widths)
    table = doc.add_table(rows=1, cols=ncols)
    _table_borders(table)

    # header
    hdr = table.rows[0]
    _repeat_header(hdr)
    heads = ["№", "ნომენკლატურა", "შტრიხ-კოდი", "ყუთი"] + (["✓"] if with_check else [])
    aligns = ["center", "left", "left", "center"] + (["center"] if with_check else [])
    for c, (t, a) in enumerate(zip(heads, aligns)):
        _cell_text(hdr.cells[c], t, align=a, bold=True, color="FFFFFF", size=10, fill=C_HEADER_FILL)

    # data
    for i, r in enumerate(rows):
        row = table.add_row()
        _cant_split(row)
        fill = C_ALT_FILL if i % 2 == 1 else None
        _cell_text(row.cells[0], i + 1, align="center", size=10, fill=fill)
        _cell_text(row.cells[1], r["item"], align="left", size=body_size, fill=fill)
        _cell_text(row.cells[2], r["barcode"], align="left", color=C_GREY, size=10, fill=fill)
        _cell_text(row.cells[3], r["boxes"], align="center", bold=True, size=body_size, fill=fill)
        if with_check:
            _cell_text(row.cells[4], check_glyph, align="center", size=14, fill=fill)

    # total row
    if total_value is not None:
        row = table.add_row()
        _cant_split(row)
        _cell_text(row.cells[0], "", fill=C_TOTAL_FILL)
        _cell_text(row.cells[1], "სულ ყუთი", align="right", bold=True, fill=C_TOTAL_FILL)
        _cell_text(row.cells[2], "", fill=C_TOTAL_FILL)
        _cell_text(row.cells[3], total_value, align="center", bold=True, size=12, fill=C_TOTAL_FILL)
        if with_check:
            _cell_text(row.cells[4], "", fill=C_TOTAL_FILL)

    _fixed_layout(table, widths)
    return table


def make_picklist_doc(veh, date_str):
    """One vehicle's pick list: per-shop breakdown with a tick column."""
    doc = Document()
    _setup_page(doc)

    _add_para(doc, [(f"შესაგროვებელი სია — {veh['vehicle']}", {"size": 15, "bold": True})],
              space_after=2)
    _add_para(doc, [(f"თარიღი: {date_str}    ·    მაღაზია: {veh['shop_count']}"
                     f"    ·    სულ ყუთი: {veh['total_boxes']}", {"size": 9, "color": C_GREY})],
              space_after=6)

    for shop in veh["shops"]:
        _add_para(doc,
                  [(shop["store"], {"size": 10, "bold": True}),
                   (f"      ·  ყუთი: {shop['boxes']}", {"size": 10, "bold": True, "color": C_GREY})],
                  space_before=6, space_after=2, fill=C_SHOP_FILL)
        _item_table(doc, shop["rows"], PICK_WIDTHS, with_check=True)
    return doc


def _delivery_note_doc(header_label, header_value, header_right, big_title,
                       sub_line, rows, total, signature,
                       receiver_label="მიმღები (სახელი/გვარი)"):
    """Shared delivery-note layout.

    The vehicle sits in the page header so it repeats on every printed page.
    """
    doc = Document()
    _setup_page(doc)
    _page_header(doc, header_label, header_value, header_right)

    _add_para(doc, [(big_title, {"size": 20, "bold": True})],
              space_before=2, space_after=3, bottom_border=True)
    _add_para(doc, [(sub_line, {"size": 9, "color": C_GREY})],
              space_before=2, space_after=6)

    _item_table(doc, rows, DEL_WIDTHS, with_check=True,
                total_value=total, body_size=11, check_glyph="☐")

    if signature:
        _add_para(doc, [("", {})], space_before=10)
        _add_para(doc, [(f"{receiver_label}: ____________________________        ",
                         {"size": 10}),
                        ("ხელმოწერა: ______________        თარიღი: ____________",
                         {"size": 10})], space_before=14, keep_next=False)
    return doc


def make_delivery_doc(note, date_str, signature=False):
    """One store's delivery note: big shop name, vehicle in the page header."""
    return _delivery_note_doc(
        header_label="მანქანა: ",
        header_value=note["vehicle"],
        header_right=f"თარიღი: {date_str}",
        big_title=note["store"],
        sub_line=f"მიღების ფურცელი  ·  სულ ყუთი: {note['boxes']}",
        rows=note["rows"],
        total=note["boxes"],
        signature=signature,
    )


def make_vehicle_delivery_doc(veh, date_str, signature=False):
    """One vehicle's consolidated delivery note (warehouse -> driver handover)."""
    return _delivery_note_doc(
        header_label="მანქანა: ",
        header_value=veh["vehicle"],
        header_right=f"მიღების ფურცელი (ჯამური)  ·  თარიღი: {date_str}",
        big_title=veh["vehicle"],
        sub_line=(f"მანქანის ჯამური დატვირთვა  ·  მაღაზია: {veh['store_count']}"
                  f"  ·  სულ ყუთი: {veh['boxes']}"),
        rows=veh["rows"],
        total=veh["boxes"],
        signature=signature,
        receiver_label="მძღოლი (სახელი/გვარი)",
    )


# ============================================================================
# 4. PACKAGING (bytes + zip)
# ============================================================================
def _doc_bytes(doc):
    b = io.BytesIO()
    doc.save(b)
    return b.getvalue()


def sanitize_filename(s, maxlen=120):
    s = re.sub(r'[\\/:*?"<>|\r\n\t]+', " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return (s[:maxlen]).strip() or "unnamed"


def _zip_bytes(files):
    """files: list of (filename, bytes). Ensures unique names."""
    buf = io.BytesIO()
    seen = {}
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        for name, data in files:
            base = name
            if name in seen:
                seen[name] += 1
                stem = name[:-5] if name.lower().endswith(".docx") else name
                base = f"{stem}_{seen[name]}.docx"
            else:
                seen[name] = 1
            z.writestr(base, data)
    return buf.getvalue()


def build_picklist_zip(df, date_str, signature=False):
    """Per vehicle: a pick list + a consolidated delivery note -> zip bytes.

    Returns (zip_bytes, n_vehicles, n_files).
    """
    veh_notes = {v["vehicle"]: v for v in build_vehicle_delivery(df)}

    files = []
    vehicles = build_pick_nested(df)
    for veh in vehicles:
        name = sanitize_filename(veh["vehicle"])
        files.append((f"{name} - შესაგროვებელი სია.docx",
                      _doc_bytes(make_picklist_doc(veh, date_str))))
        note = veh_notes.get(veh["vehicle"])
        if note is not None:
            files.append((f"{name} - მიღების ფურცელი.docx",
                          _doc_bytes(make_vehicle_delivery_doc(note, date_str, signature))))
    return _zip_bytes(files), len(vehicles), len(files)


def build_delivery_zip(df, date_str, signature=False, group_by_vehicle=True,
                       include_picklist=True):
    """Store delivery notes -> zip bytes, with each vehicle's pick list included.

    group_by_vehicle=True nests each store's note inside a folder named after
    its vehicle, and (when include_picklist) drops that vehicle's pick list in
    the same folder so each folder is a complete package for one truck:
        "<vehicle>/00 - შესაგროვებელი სია.docx"
        "<vehicle>/<store>.docx"
    group_by_vehicle=False produces a flat list named "<vehicle> - <store>.docx".

    Returns (zip_bytes, n_store_notes, n_files).
    """
    files = []

    # one pick list per vehicle, first in its folder ("00 - " sorts it to the top)
    if include_picklist:
        for veh in build_pick_nested(df):
            name = sanitize_filename(veh["vehicle"])
            path = (f"{name}/00 - შესაგროვებელი სია.docx" if group_by_vehicle
                    else f"{name} - 00 - შესაგროვებელი სია.docx")
            files.append((path, _doc_bytes(make_picklist_doc(veh, date_str))))

    n_notes = 0
    for note in build_delivery_flat(df):
        if group_by_vehicle:
            path = (f"{sanitize_filename(note['vehicle'])}/"
                    f"{sanitize_filename(note['store'], 100)}.docx")
        else:
            path = (f"{sanitize_filename(note['vehicle'])} - "
                    f"{sanitize_filename(note['store'], 90)}.docx")
        files.append((path, _doc_bytes(make_delivery_doc(note, date_str, signature))))
        n_notes += 1

    return _zip_bytes(files), n_notes, len(files)
